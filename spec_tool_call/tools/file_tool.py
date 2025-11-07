"""Enhanced file reading tool supporting diverse file types."""
import os
import mimetypes
from typing import Dict, Any, Optional
from pathlib import Path

import pandas as pd


def read_file(path: str, extract_text: bool = True) -> Dict[str, Any]:
    """
    Read various file types and return structured information.
    
    Supported formats:
    - Text: .txt, .md, .json, .xml, .html, .csv
    - Spreadsheets: .xlsx, .xls, .csv
    - Documents: .pdf, .docx
    - Data: .json, .jsonl, .yaml, .yml
    
    Args:
        path: Path to the file
        extract_text: Whether to extract full text content
        
    Returns:
        Dict with file info, type, and content/preview
    """
    if not os.path.exists(path):
        return {
            "error": "file_not_found",
            "path": path,
            "message": f"File does not exist: {path}"
        }
    
    # Get file info
    file_stat = os.stat(path)
    file_size = file_stat.st_size
    file_ext = Path(path).suffix.lower()
    
    # Detect mime type
    mime_type, _ = mimetypes.guess_type(path)
    
    base_info = {
        "path": path,
        "size_bytes": file_size,
        "extension": file_ext,
        "mime_type": mime_type,
    }
    
    # Route to appropriate handler based on extension
    try:
        if file_ext in ('.txt', '.md', '.log'):
            return {**base_info, **_read_text_file(path, extract_text)}
        
        elif file_ext in ('.csv',):
            return {**base_info, **_read_csv_file(path)}
        
        elif file_ext in ('.xlsx', '.xls'):
            return {**base_info, **_read_excel_file(path)}
        
        elif file_ext in ('.json',):
            return {**base_info, **_read_json_file(path)}
        
        elif file_ext in ('.jsonl',):
            return {**base_info, **_read_jsonl_file(path)}
        
        elif file_ext in ('.xml', '.html', '.htm'):
            return {**base_info, **_read_markup_file(path)}
        
        elif file_ext in ('.pdf',):
            return {**base_info, **_read_pdf_file(path, extract_text)}
        
        elif file_ext in ('.docx',):
            return {**base_info, **_read_docx_file(path)}
        
        elif file_ext in ('.yaml', '.yml'):
            return {**base_info, **_read_yaml_file(path)}
        
        else:
            # Try to read as text
            try:
                return {**base_info, **_read_text_file(path, extract_text)}
            except UnicodeDecodeError:
                return {
                    **base_info,
                    "kind": "binary",
                    "message": "Binary file - cannot extract text",
                    "preview": f"Binary file of size {file_size} bytes"
                }
    
    except Exception as e:
        return {
            **base_info,
            "error": "read_error",
            "message": f"Error reading file: {str(e)}"
        }


# ============================================================================
# File type handlers
# ============================================================================

def _read_text_file(path: str, full_text: bool = True) -> Dict[str, Any]:
    """Read plain text file."""
    with open(path, 'r', encoding='utf-8', errors='ignore') as f:
        content = f.read()
    
    lines = content.split('\n')
    
    return {
        "kind": "text",
        "line_count": len(lines),
        "char_count": len(content),
        "preview": content[:2000] if not full_text else content,
        "full_text": content if full_text else None
    }


def _read_csv_file(path: str) -> Dict[str, Any]:
    """Read CSV file."""
    df = pd.read_csv(path)
    
    return {
        "kind": "csv",
        "rows": len(df),
        "columns": len(df.columns),
        "column_names": df.columns.tolist(),
        "preview_rows": df.head(10).to_dict(orient='records'),
        "dtypes": df.dtypes.astype(str).to_dict()
    }


def _read_excel_file(path: str) -> Dict[str, Any]:
    """Read Excel file."""
    # Read all sheets
    excel_file = pd.ExcelFile(path)
    sheet_names = excel_file.sheet_names
    
    # Read first sheet for preview
    df = pd.read_excel(path, sheet_name=sheet_names[0])
    
    return {
        "kind": "excel",
        "sheets": sheet_names,
        "sheet_count": len(sheet_names),
        "active_sheet": sheet_names[0],
        "rows": len(df),
        "columns": len(df.columns),
        "column_names": df.columns.tolist(),
        "preview_rows": df.head(10).to_dict(orient='records')
    }


def _read_json_file(path: str) -> Dict[str, Any]:
    """Read JSON file."""
    import json
    
    with open(path, 'r', encoding='utf-8') as f:
        data = json.load(f)
    
    return {
        "kind": "json",
        "data_type": type(data).__name__,
        "size": len(data) if isinstance(data, (list, dict)) else None,
        "preview": str(data)[:2000],
        "data": data
    }


def _read_jsonl_file(path: str) -> Dict[str, Any]:
    """Read JSON Lines file."""
    import json
    
    lines = []
    with open(path, 'r', encoding='utf-8') as f:
        for line in f:
            if line.strip():
                lines.append(json.loads(line))
    
    return {
        "kind": "jsonl",
        "line_count": len(lines),
        "preview": lines[:10],
        "data": lines
    }


def _read_markup_file(path: str) -> Dict[str, Any]:
    """Read XML/HTML file."""
    with open(path, 'r', encoding='utf-8', errors='ignore') as f:
        content = f.read()
    
    # Try to extract text from HTML
    try:
        from bs4 import BeautifulSoup
        soup = BeautifulSoup(content, 'html.parser')
        text = soup.get_text(separator=' ', strip=True)
        
        return {
            "kind": "markup",
            "format": "html" if path.endswith(('.html', '.htm')) else "xml",
            "raw_size": len(content),
            "text_content": text[:2000],
            "preview": content[:1000]
        }
    except ImportError:
        return {
            "kind": "markup",
            "format": "html" if path.endswith(('.html', '.htm')) else "xml",
            "preview": content[:2000],
            "message": "Install beautifulsoup4 for better parsing"
        }


def _read_pdf_file(path: str, extract_text: bool = True) -> Dict[str, Any]:
    """Read PDF file."""
    try:
        import pdfplumber
        
        with pdfplumber.open(path) as pdf:
            page_count = len(pdf.pages)
            
            if extract_text:
                text_parts = []
                for i, page in enumerate(pdf.pages[:5]):  # First 5 pages
                    text = page.extract_text()
                    if text:
                        text_parts.append(f"--- Page {i+1} ---\n{text}")
                
                full_text = "\n\n".join(text_parts)
                
                return {
                    "kind": "pdf",
                    "pages": page_count,
                    "text_preview": full_text[:3000],
                    "extracted_pages": min(5, page_count)
                }
            else:
                return {
                    "kind": "pdf",
                    "pages": page_count,
                    "message": "Text extraction disabled"
                }
    
    except ImportError:
        return {
            "kind": "pdf",
            "message": "Install pdfplumber for PDF text extraction",
            "preview": f"PDF file - {os.path.getsize(path)} bytes"
        }
    except Exception as e:
        return {
            "kind": "pdf",
            "error": f"PDF extraction error: {str(e)}",
            "preview": f"PDF file - {os.path.getsize(path)} bytes"
        }


def _read_docx_file(path: str) -> Dict[str, Any]:
    """Read DOCX file."""
    try:
        from docx import Document
        
        doc = Document(path)
        
        # Extract text from paragraphs
        text_parts = [para.text for para in doc.paragraphs if para.text.strip()]
        full_text = "\n\n".join(text_parts)
        
        return {
            "kind": "docx",
            "paragraphs": len(doc.paragraphs),
            "text_preview": full_text[:3000],
            "char_count": len(full_text)
        }
    
    except ImportError:
        return {
            "kind": "docx",
            "message": "Install python-docx for DOCX reading",
            "preview": f"DOCX file - {os.path.getsize(path)} bytes"
        }
    except Exception as e:
        return {
            "kind": "docx",
            "error": f"DOCX extraction error: {str(e)}"
        }


def _read_yaml_file(path: str) -> Dict[str, Any]:
    """Read YAML file."""
    try:
        import yaml
        
        with open(path, 'r', encoding='utf-8') as f:
            data = yaml.safe_load(f)
        
        return {
            "kind": "yaml",
            "data_type": type(data).__name__,
            "preview": str(data)[:2000],
            "data": data
        }
    
    except ImportError:
        return {
            "kind": "yaml",
            "message": "Install pyyaml for YAML parsing",
            "preview": open(path, 'r').read()[:1000]
        }


if __name__ == "__main__":
    """Test file reading capabilities"""
    print("=" * 70)
    print("FILE READING TOOL TEST")
    print("=" * 70)
    
    # Test with a text file
    test_file = __file__  # Read itself
    print(f"\n[TEST] Reading Python file: {test_file}")
    result = read_file(test_file)
    print(f"Result: {result.get('kind')}, {result.get('line_count')} lines")
    print(f"Preview: {result.get('preview')[:200]}...")

